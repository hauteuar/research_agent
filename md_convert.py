from docx import Document
from markdown import markdown as md_to_html
from bs4 import BeautifulSoup
import re

def clean_markdown(md):
    md = re.sub(r'(?m)^#+\s*', '', md)  # Remove all heading #
    md = re.sub(r'\*+', '', md)        # Remove asterisks
    return md

# Load your full markdown content here
with open("input.md", "r", encoding="utf-8") as f:
    markdown_text = f.read()

cleaned_md = clean_markdown(markdown_text)
html = md_to_html(cleaned_md)
soup = BeautifulSoup(html, "html.parser")

doc = Document()
for elem in soup.descendants:
    if elem.name == "p":
        doc.add_paragraph(elem.get_text())
    elif elem.name == "li":
        doc.add_paragraph(f"- {elem.get_text()}")
    elif elem.name == "h1":
        doc.add_heading(elem.get_text(), level=1)
    elif elem.name == "h2":
        doc.add_heading(elem.get_text(), level=2)
    elif elem.name == "h3":
        doc.add_heading(elem.get_text(), level=3)
    elif elem.name == "pre":
        doc.add_paragraph(elem.get_text())

doc.save("Opulence_Flow_Converted.docx")
